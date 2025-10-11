import random
import json
import os
import uuid
from datetime import datetime
from docx import Document
from typing import List, Dict, Any, Optional

from .models.llm_loader import QuizGenerator, QuizEvaluator
from .document_processing import extract_text

# Initialize models
quiz_generator = QuizGenerator()
quiz_evaluator = QuizEvaluator()

# Directory for quiz export
QUIZ_DIR = os.path.join(os.getcwd(), "data", "quizzes")
os.makedirs(QUIZ_DIR, exist_ok=True)


def generate_quiz_from_prompt(topic: str, num_questions: int, question_type: str, 
                            difficulty: str, title: Optional[str] = None) -> Dict[str, Any]:
    """Generate a quiz from a text prompt or topic."""
    print(f"[INFO] Generating quiz from topic: {topic}")
    
    quiz_data = quiz_generator.generate_quiz_from_topic(
        topic=topic,
        num_questions=num_questions,
        question_type=question_type,
        difficulty=difficulty
    )
    
    # Override title if provided
    if title:
        quiz_data["title"] = title
    
    # Add unique ID and metadata
    quiz_data["id"] = str(uuid.uuid4())
    quiz_data["created_at"] = datetime.utcnow().isoformat()
    
    return quiz_data


def generate_quiz_from_document(file_path: str, num_questions: int, 
                              question_type: str, difficulty: str,
                              title: Optional[str] = None) -> Dict[str, Any]:
    """Generate a quiz from an uploaded document."""
    print(f"[INFO] Generating quiz from document: {file_path}")
    
    # Extract text from document
    content_text = extract_text(file_path)
    
    if not content_text.strip():
        raise ValueError("No text content found in the document")
    
    quiz_data = quiz_generator.generate_quiz_from_document(
        content=content_text,
        num_questions=num_questions,
        question_type=question_type,
        difficulty=difficulty
    )
    
    # Override title if provided
    if title:
        quiz_data["title"] = title
    elif not quiz_data.get("title"):
        quiz_data["title"] = f"Quiz from {os.path.basename(file_path)}"
    
    # Add unique ID and metadata
    quiz_data["id"] = str(uuid.uuid4())
    quiz_data["created_at"] = datetime.utcnow().isoformat()
    quiz_data["source_document"] = os.path.basename(file_path)
    
    return quiz_data


def generate_quiz_from_mistakes(student_id: str, mistakes: List[Dict], 
                              num_questions: int, question_type: str, 
                              difficulty: str, title: Optional[str] = None) -> Dict[str, Any]:
    """Generate a quiz based on a student's previous mistakes."""
    print(f"[INFO] Generating quiz for student {student_id} based on mistakes...")
    
    if not mistakes:
        # If no mistakes, generate a general review quiz
        return generate_quiz_from_prompt(
            topic="General Review",
            num_questions=num_questions,
            question_type=question_type,
            difficulty=difficulty,
            title=title or "Practice Quiz"
        )
    
    quiz_data = quiz_generator.generate_quiz_from_mistakes(
        mistakes=mistakes,
        num_questions=num_questions,
        question_type=question_type,
        difficulty=difficulty
    )
    
    # Override title if provided
    if title:
        quiz_data["title"] = title
    
    # Add unique ID and metadata
    quiz_data["id"] = str(uuid.uuid4())
    quiz_data["created_at"] = datetime.utcnow().isoformat()
    quiz_data["target_student"] = student_id
    
    return quiz_data


def evaluate_quiz_submission(quiz_data: Dict[str, Any], 
                           student_responses: Dict[str, str],
                           time_taken: int = 0) -> Dict[str, Any]:
    """Evaluate a student's quiz submission."""
    print("[INFO] Evaluating quiz submission...")
    
    questions = quiz_data.get("questions", [])
    question_type = quiz_data.get("question_type", "MCQ")
    
    total_points = 0
    earned_points = 0
    detailed_results = []
    incorrect_responses = []
    
    for question in questions:
        question_id = question.get("id")
        student_answer = student_responses.get(question_id, "").strip()
        
        # Evaluate the response
        evaluation = quiz_evaluator.evaluate_response(
            question=question,
            student_answer=student_answer,
            question_type=question_type
        )
        
        total_points += evaluation["max_points"]
        earned_points += evaluation["points_earned"]
        
        # Prepare detailed result
        result = {
            "question_id": question_id,
            "question_text": question.get("question", ""),
            "student_answer": student_answer,
            "correct_answer": question.get("correct_answer", ""),
            "is_correct": evaluation["is_correct"],
            "points_earned": evaluation["points_earned"],
            "max_points": evaluation["max_points"],
            "feedback": evaluation.get("feedback", ""),
            "response_time": 0  # Could be tracked per question in future
        }
        
        detailed_results.append(result)
        
        # Track incorrect responses for mistake logging
        if not evaluation["is_correct"]:
            incorrect_responses.append({
                "question": question.get("question", ""),
                "topic": quiz_data.get("topic", "General"),
                "concept": _extract_concept_from_question(question),
                "mistake_type": _classify_mistake_type(question, student_answer),
                "student_answer": student_answer,
                "correct_answer": question.get("correct_answer", "")
            })
    
    # Calculate final score
    percentage = (earned_points / total_points * 100) if total_points > 0 else 0
    
    return {
        "score": earned_points,
        "max_score": total_points,
        "percentage": percentage,
        "correct_answers": sum(1 for r in detailed_results if r["is_correct"]),
        "total_questions": len(questions),
        "time_taken": time_taken,
        "detailed_results": detailed_results,
        "incorrect_responses": incorrect_responses,
        "performance_level": _get_performance_level(percentage)
    }


def update_quiz_questions(quiz_data: Dict[str, Any], 
                         updated_questions: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Update quiz questions (for teacher editing)."""
    print("[INFO] Updating quiz questions...")
    
    # Create a mapping of question IDs to questions
    question_map = {q.get("id"): q for q in quiz_data.get("questions", [])}
    
    # Update existing questions
    for updated_question in updated_questions:
        question_id = updated_question.get("id")
        if question_id in question_map:
            question_map[question_id].update(updated_question)
    
    # Update the quiz data
    quiz_data["questions"] = list(question_map.values())
    quiz_data["updated_at"] = datetime.utcnow().isoformat()
    quiz_data["num_questions"] = len(quiz_data["questions"])
    
    return quiz_data


def export_quiz_to_docx(quiz_data: Dict[str, Any], include_answers: bool = True) -> str:
    """Export quiz to DOCX file."""
    print("[INFO] Exporting quiz to DOCX...")
    
    doc = Document()
    
    # Add title and metadata
    doc.add_heading(quiz_data.get("title", "Quiz"), level=1)
    
    info_paragraph = doc.add_paragraph()
    info_paragraph.add_run("Topic: ").bold = True
    info_paragraph.add_run(f"{quiz_data.get('topic', 'General')}\n")
    info_paragraph.add_run("Difficulty: ").bold = True
    info_paragraph.add_run(f"{quiz_data.get('difficulty', 'Medium')}\n")
    info_paragraph.add_run("Questions: ").bold = True
    info_paragraph.add_run(f"{quiz_data.get('num_questions', 0)}\n")
    info_paragraph.add_run("Type: ").bold = True
    info_paragraph.add_run(f"{quiz_data.get('question_type', 'MCQ')}")
    
    doc.add_paragraph()  # Empty line
    
    # Add questions
    for i, question in enumerate(quiz_data.get("questions", []), 1):
        # Question text
        question_paragraph = doc.add_paragraph()
        question_paragraph.add_run(f"Question {i}: ").bold = True
        question_paragraph.add_run(question.get("question", ""))
        
        # Options for MCQ
        if quiz_data.get("question_type") == "MCQ" and question.get("options"):
            for j, option in enumerate(question["options"]):
                option_letter = chr(65 + j)  # A, B, C, D
                doc.add_paragraph(f"   {option_letter}) {option}")
        
        # Answer and explanation (if including answers)
        if include_answers:
            answer_paragraph = doc.add_paragraph()
            answer_paragraph.add_run("Answer: ").bold = True
            answer_paragraph.add_run(question.get("correct_answer", ""))
            
            if question.get("explanation"):
                explanation_paragraph = doc.add_paragraph()
                explanation_paragraph.add_run("Explanation: ").bold = True
                explanation_paragraph.add_run(question.get("explanation", ""))
        
        doc.add_paragraph()  # Empty line between questions
    
    # Save file
    filename = f"quiz_{quiz_data.get('id', 'unknown')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join(QUIZ_DIR, filename)
    doc.save(file_path)
    
    return file_path


def export_quiz_to_json(quiz_data: Dict[str, Any]) -> str:
    """Export quiz to JSON file."""
    print("[INFO] Exporting quiz to JSON...")
    
    filename = f"quiz_{quiz_data.get('id', 'unknown')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    file_path = os.path.join(QUIZ_DIR, filename)
    
    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(quiz_data, f, indent=2, ensure_ascii=False)
    
    return file_path


def create_quiz_share_code() -> str:
    """Generate a unique share code for quiz sharing."""
    return str(uuid.uuid4())[:8].upper()


def _extract_concept_from_question(question: Dict[str, Any]) -> str:
    """Extract the main concept from a question for mistake tracking."""
    # This is a simplified implementation
    # In a real system, you might use NLP to extract key concepts
    question_text = question.get("question", "").lower()
    
    # Simple keyword extraction (can be enhanced with NLP)
    if "calculate" in question_text or "compute" in question_text:
        return "Calculation"
    elif "define" in question_text or "what is" in question_text:
        return "Definition"
    elif "explain" in question_text or "describe" in question_text:
        return "Explanation"
    elif "compare" in question_text or "contrast" in question_text:
        return "Comparison"
    else:
        return "General Concept"


def _classify_mistake_type(question: Dict[str, Any], student_answer: str) -> str:
    """Classify the type of mistake made by the student."""
    # This is a simplified classification
    # In a real system, you might use more sophisticated analysis
    
    if not student_answer.strip():
        return "no_response"
    
    question_type = question.get("type", "MCQ")
    
    if question_type == "MCQ":
        # For MCQ, it's usually a conceptual mistake
        return "conceptual"
    elif question_type in ["Short Answer", "Essay"]:
        # Analyze the length and content of the response
        if len(student_answer.split()) < 3:
            return "incomplete"
        else:
            return "conceptual"
    else:
        return "conceptual"


def _get_performance_level(percentage: float) -> str:
    """Get performance level based on percentage score."""
    if percentage >= 90:
        return "Excellent"
    elif percentage >= 80:
        return "Good"
    elif percentage >= 70:
        return "Satisfactory"
    elif percentage >= 60:
        return "Needs Improvement"
    else:
        return "Poor"


# Utility functions for document processing
def chunk_text(text: str, max_length: int = 2000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks for large document processing."""
    words = text.split()
    chunks = []
    start = 0
    
    while start < len(words):
        end = start + max_length
        chunk = " ".join(words[start:end])
        chunks.append(chunk)
        start += max_length - overlap
        
        if start >= len(words):
            break
    
    return chunks